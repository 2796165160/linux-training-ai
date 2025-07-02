# backend/app/utils/docx_helpers.py

from docx import Document
import re
import io
import logging

logger = logging.getLogger(__name__)

def extract_placeholders(doc):
    """从DOCX文档中提取占位符"""
    placeholders = set()
    pattern = r"\{([^{}]+)\}"
    
    # 检查段落
    for para in doc.paragraphs:
        matches = re.findall(pattern, para.text)
        placeholders.update(matches)
    
    # 检查表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    matches = re.findall(pattern, para.text)
                    placeholders.update(matches)
    
    return list(placeholders)

def apply_docx_template(docx_content, report_sections):
    """将报告内容应用到DOCX模板"""
    try:
        # 加载模板
        doc = Document(io.BytesIO(docx_content))
        
        # 提取模板中的占位符
        placeholders = extract_placeholders(doc)
        logger.info(f"提取到的占位符: {placeholders}")
        
        # 替换段落中的占位符
        for para in doc.paragraphs:
            for key in report_sections:
                placeholder = f"{{{key}}}"
                if placeholder in para.text:
                    para.text = para.text.replace(placeholder, report_sections[key])
        
        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key in report_sections:
                            placeholder = f"{{{key}}}"
                            if placeholder in para.text:
                                para.text = para.text.replace(placeholder, report_sections[key])
        
        # 保存到内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output
    except Exception as e:
        logger.error(f"应用模板时出错: {str(e)}")
        raise
