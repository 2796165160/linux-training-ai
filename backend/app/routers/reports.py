# from fastapi import APIRouter, Depends, HTTPException, status, Response
# from sqlalchemy.orm import Session
# from typing import List, Optional
# import io
# from docx import Document
# from reportlab.pdfgen import canvas
# from reportlab.lib.pagesizes import letter

# from ..database import get_db
# from ..models.user import User
# from ..models.report import Report
# from ..schemas.report import ReportCreate, ReportUpdate, ReportResponse
# from ..utils.security import get_current_user
# from ..services.ai_service import generate_report_with_qwen
# from ..config import settings

# router = APIRouter(prefix=f"{settings.API_V1_STR}/reports", tags=["报告"])

# @router.post("/generate", response_model=ReportResponse)
# async def create_report(
#     report_data: ReportCreate,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     """生成新报告"""
#     report = await generate_report_with_qwen(
#         db=db, 
#         task_id=report_data.task_id, 
#         user_id=current_user.id,
#         template_id=report_data.template_id
#     )
    
#     if not report:
#         raise HTTPException(status_code=500, detail="报告生成失败")
    
#     return report

# @router.get("/", response_model=List[ReportResponse])
# async def read_reports(
#     skip: int = 0,
#     limit: int = 100,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     """获取当前用户的报告列表"""
#     reports = db.query(Report).filter(Report.user_id == current_user.id).offset(skip).limit(limit).all()
#     return reports

# @router.get("/{report_id}", response_model=ReportResponse)
# async def read_report(
#     report_id: int,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     """获取指定报告详情"""
#     report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
#     if report is None:
#         raise HTTPException(status_code=404, detail="报告不存在")
#     return report

# @router.put("/{report_id}", response_model=ReportResponse)
# async def update_report(
#     report_id: int,
#     report_data: ReportUpdate,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     """更新报告内容"""
#     report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
#     if report is None:
#         raise HTTPException(status_code=404, detail="报告不存在")
    
#     report.content = report_data.content
#     db.commit()
#     db.refresh(report)
    
#     return report

# @router.delete("/{report_id}", status_code=status.HTTP_204_NO_CONTENT)
# async def delete_report(
#     report_id: int,
#     db: Session = Depends(get_db),
#     current_user: User = Depends(get_current_user)
# ):
#     """删除报告"""
#     report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
#     if report is None:
#         raise HTTPException(status_code=404, detail="报告不存在")
    
#     db.delete(report)
#     db.commit()
#     return None

# @router.get("/{report_id}/export/docx")
# async def export_report_docx(
#     report_id: int,
#     db: Session = Depends(get_db)
# ):
#     """导出报告为Word格式 (无需认证)"""
#     try:
#         # 查找报告
#         report = db.query(Report).filter(Report.id == report_id).first()
#         if report is None:
#             raise HTTPException(status_code=404, detail="报告不存在")
        
#         from docx import Document
#         import io
        
#         # 创建文档
#         doc = Document()
#         doc.add_heading(report.title, 0)
        
#         # 添加内容 - 简单地作为纯文本添加
#         for paragraph in report.content.split('\n\n'):
#             if paragraph.strip():
#                 doc.add_paragraph(paragraph)
        
#         # 保存到内存
#         buffer = io.BytesIO()
#         doc.save(buffer)
#         buffer.getvalue()  # 确保缓冲区已写入
#         buffer.seek(0)
        
#         # 使用安全的ASCII文件名 - 关键修复!
#         safe_filename = f"report_{report_id}.docx"
        
#         # 返回文件
#         from fastapi.responses import Response
#         return Response(
#             content=buffer.getvalue(),
#             media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#             headers={"Content-Disposition": f"attachment; filename={safe_filename}"}
#         )
#     except Exception as e:
#         import traceback
#         error_msg = traceback.format_exc()
#         print(f"Word导出错误: {error_msg}")
#         raise HTTPException(status_code=500, detail="Word文档生成失败，请稍后再试")

# @router.get("/{report_id}/export/html")
# async def export_report_html(
#     report_id: int,
#     db: Session = Depends(get_db)
# ):
#     """导出报告为HTML格式(可以在浏览器中打印为PDF)"""
#     try:
#         # 查找报告
#         report = db.query(Report).filter(Report.id == report_id).first()
#         if report is None:
#             raise HTTPException(status_code=404, detail="报告不存在")
        
#         # 转义HTML特殊字符
#         import html
#         title = html.escape(report.title)
#         content = html.escape(report.content)
#         content = content.replace('\n', '<br>')  # 保留换行
        
#         # 创建HTML内容
#         html_content = f"""
#         <!DOCTYPE html>
#         <html lang="zh-CN">
#         <head>
#             <meta charset="UTF-8">
#             <meta name="viewport" content="width=device-width, initial-scale=1.0">
#             <title>{title}</title>
#             <style>
#                 @media print {{
#                     body {{ margin: 1cm; }}
#                 }}
#                 body {{
#                     font-family: "Microsoft YaHei", Arial, sans-serif;
#                     line-height: 1.6;
#                     margin: 2cm;
#                     color: #333;
#                 }}
#                 h1 {{ 
#                     text-align: center;
#                     margin-bottom: 20px;
#                     color: #000;
#                 }}
#                 .content {{ 
#                     text-align: justify;
#                 }}
#                 .print-button {{
#                     text-align: center;
#                     margin: 20px 0;
#                 }}
#                 @page {{ 
#                     size: A4; 
#                     margin: 2cm;
#                 }}
#                 /* 打印时隐藏按钮 */
#                 @media print {{
#                     .print-button {{ display: none; }}
#                 }}
#             </style>
#         </head>
#         <body>
#             <div class="print-button">
#                 <button onclick="window.print()" style="
#                     background-color: #4CAF50;
#                     color: white;
#                     padding: 10px 20px;
#                     border: none;
#                     border-radius: 5px;
#                     cursor: pointer;
#                     font-size: 16px;">
#                     打印为PDF
#                 </button>
#             </div>
#             <h1>{title}</h1>
#             <div class="content">{content}</div>
#             <div class="print-button">
#                 <button onclick="window.print()" style="
#                     background-color: #4CAF50;
#                     color: white;
#                     padding: 10px 20px;
#                     border: none;
#                     border-radius: 5px;
#                     cursor: pointer;
#                     font-size: 16px;">
#                     打印为PDF
#                 </button>
#             </div>
#         </body>
#         </html>
#         """
        
#         # 返回HTML - 注意这里不是附件下载而是直接在浏览器中显示
#         from fastapi.responses import HTMLResponse
#         return HTMLResponse(content=html_content)
#     except Exception as e:
#         import traceback
#         error_msg = traceback.format_exc()
#         print(f"HTML导出错误: {error_msg}")
#         raise HTTPException(status_code=500, detail="HTML导出失败")


from fastapi import APIRouter, Depends, HTTPException, status, Response
from sqlalchemy.orm import Session
from typing import List, Optional
import io
import time
import logging
from datetime import datetime
from fastapi.responses import StreamingResponse, HTMLResponse

from ..database import get_db
from ..models.user import User
from ..models.report import Report
from ..schemas.report import ReportCreate, ReportUpdate, ReportResponse
from ..utils.security import get_current_user
from ..services.ai_service import generate_report_with_qwen
from ..config import settings

router = APIRouter(prefix=f"{settings.API_V1_STR}/reports", tags=["报告"])

# 配置日志
logger = logging.getLogger(__name__)

@router.post("/generate", response_model=ReportResponse)
async def create_report(
    report_data: ReportCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """生成新报告"""
    report = await generate_report_with_qwen(
        db=db, 
        task_id=report_data.task_id, 
        user_id=current_user.id,
        template_id=report_data.template_id
    )
    
    if not report:
        raise HTTPException(status_code=500, detail="报告生成失败")
    
    return report

@router.get("/", response_model=List[ReportResponse])
async def read_reports(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取当前用户的报告列表"""
    reports = db.query(Report).filter(Report.user_id == current_user.id).offset(skip).limit(limit).all()
    return reports

@router.get("/{report_id}", response_model=ReportResponse)
async def read_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取指定报告详情"""
    report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
    if report is None:
        raise HTTPException(status_code=404, detail="报告不存在")
    return report

@router.put("/{report_id}", response_model=ReportResponse)
async def update_report(
    report_id: int,
    report_data: ReportUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """更新报告内容"""
    report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
    if report is None:
        raise HTTPException(status_code=404, detail="报告不存在")
    
    report.content = report_data.content
    db.commit()
    db.refresh(report)
    
    return report

@router.delete("/{report_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_report(
    report_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """删除报告"""
    report = db.query(Report).filter(Report.id == report_id, Report.user_id == current_user.id).first()
    if report is None:
        raise HTTPException(status_code=404, detail="报告不存在")
    
    db.delete(report)
    db.commit()
    return None

@router.get("/{report_id}/export/docx")
async def export_report_docx(
    report_id: int,
    db: Session = Depends(get_db)
):
    """导出报告为Word格式"""
    try:
        # 查找报告
        report = db.query(Report).filter(Report.id == report_id).first()
        if report is None:
            raise HTTPException(status_code=404, detail="报告不存在")
        
        from docx import Document
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading(report.title, 0)
        
        # 添加内容 - 解析Markdown格式
        paragraphs = report.content.split('\n')
        for p in paragraphs:
            if p.strip():
                if p.startswith('# '):
                    doc.add_heading(p[2:], 1)
                elif p.startswith('## '):
                    doc.add_heading(p[3:], 2)
                elif p.startswith('### '):
                    doc.add_heading(p[4:], 3)
                elif p.startswith('- '):
                    # 列表项
                    doc.add_paragraph(p[2:], style='ListBullet')
                elif p.startswith('1. ') or p.startswith('2. ') or p.startswith('3. '):
                    # 有序列表
                    doc.add_paragraph(p[3:], style='ListNumber')
                else:
                    doc.add_paragraph(p)
        
        # 保存到内存
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        # 安全的文件名
        safe_filename = f"report_{report_id}.docx"
        
        # 返回文件
        return Response(
            content=f.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={safe_filename}"}
        )
    except Exception as e:
        logger.error(f"Word导出错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Word文档生成失败: {str(e)}")

@router.get("/{report_id}/export/pdf")
async def export_report_pdf(
    report_id: int,
    db: Session = Depends(get_db)
):
    """导出报告为PDF格式"""
    try:
        # 查找报告
        report = db.query(Report).filter(Report.id == report_id).first()
        if report is None:
            raise HTTPException(status_code=404, detail="报告不存在")
        
        # 使用更可靠的方法生成PDF
        try:
            # 尝试使用FPDF2生成PDF (对中文支持更好)
            from fpdf import FPDF
            
            # 创建PDF
            pdf = FPDF()
            pdf.add_page()
            
            # 添加标题
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, "AI实训报告", ln=True, align='C')
            pdf.ln(5)
            
            # 添加报告标题
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, report.title, ln=True)
            pdf.ln(5)
            
            # 添加内容 (简化处理，将所有内容作为文本添加)
            pdf.set_font("Arial", "", 12)
            
            # 分段落处理内容
            paragraphs = report.content.split("\n\n")
            for paragraph in paragraphs:
                if paragraph.strip():
                    # 添加段落
                    pdf.multi_cell(0, 10, paragraph.strip())
                    pdf.ln(5)
            
            # 保存到内存缓冲区
            buffer = io.BytesIO()
            pdf.output(buffer)
            buffer.seek(0)
            
            # 使用StreamingResponse返回PDF
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=report_{report_id}.pdf"}
            )
            
        except ImportError:
            # 如果FPDF不可用，使用ReportLab
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            
            # 创建PDF文件
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            
            # 定义自定义样式
            styles.add(ParagraphStyle(
                name='CenterTitle',
                parent=styles['Heading1'],
                alignment=1,  # 居中
            ))
            
            # 创建内容列表
            content = []
            
            # 添加标题
            content.append(Paragraph("AI实训报告", styles['CenterTitle']))
            content.append(Spacer(1, 12))
            content.append(Paragraph(report.title, styles['Heading1']))
            content.append(Spacer(1, 12))
            
            # 添加报告内容
            paragraphs = report.content.split("\n\n")
            for paragraph in paragraphs:
                if paragraph.strip():
                    content.append(Paragraph(paragraph.strip(), styles['Normal']))
                    content.append(Spacer(1, 6))
            
            # 构建PDF
            doc.build(content)
            buffer.seek(0)
            
            # 使用StreamingResponse返回PDF
            return StreamingResponse(
                buffer,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=report_{report_id}.pdf"}
            )
        
    except Exception as e:
        logger.error(f"PDF生成错误: {str(e)}", exc_info=True)
        # 尝试提供友好的错误信息
        raise HTTPException(
            status_code=500,
            detail="PDF生成失败，可能是因为报告内容包含不支持的字符。请尝试导出为HTML或文本格式。"
        )

@router.get("/{report_id}/export/html")
async def export_report_html(
    report_id: int,
    db: Session = Depends(get_db)
):
    """导出报告为HTML格式(可在浏览器中打印为PDF)"""
    try:
        # 查找报告
        report = db.query(Report).filter(Report.id == report_id).first()
        if report is None:
            raise HTTPException(status_code=404, detail="报告不存在")
        
        # 处理HTML特殊字符
        import html
        title_safe = html.escape(report.title)
        
        # Markdown转HTML（简易版本）
        import re
        content = report.content
        
        # 替换标题
        content = re.sub(r'^# (.*?)$', r'<h1>\1</h1>', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.*?)$', r'<h2>\1</h2>', content, flags=re.MULTILINE)
        content = re.sub(r'^### (.*?)$', r'<h3>\1</h3>', content, flags=re.MULTILINE)
        
        # 替换列表
        content = re.sub(r'^- (.*?)$', r'<li>\1</li>', content, flags=re.MULTILINE)
        content = re.sub(r'<li>.*?</li>(\n<li>.*?</li>)+', lambda m: f'<ul>{m.group(0)}</ul>', content, flags=re.DOTALL)
        
        # 替换粗体和斜体
        content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', content)
        content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', content)
        
        # 替换段落和换行
        content = re.sub(r'(?<!\n)\n(?!\n)', r'<br>', content)
        content = re.sub(r'\n\n+', r'</p><p>', content)
        content = f'<p>{content}</p>'
        
        # 创建HTML内容
        html_content = f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{title_safe}</title>
            <style>
                @media print {{
                    body {{ 
                        margin: 2cm; 
                        font-family: "SimSun", "宋体", serif;
                    }}
                    .no-print {{ display: none; }}
                }}
                body {{
                    font-family: "Microsoft YaHei", "微软雅黑", "SimSun", "宋体", sans-serif;
                    line-height: 1.6;
                    margin: 2cm;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 2cm;
                }}
                h1, h2, h3, h4 {{ 
                    font-weight: bold;
                    margin-top: 1em;
                    margin-bottom: 0.5em;
                    color: #000;
                }}
                h1 {{ font-size: 24pt; }}
                h2 {{ font-size: 18pt; }}
                h3 {{ font-size: 14pt; }}
                p {{ 
                    margin-bottom: 0.5em; 
                    text-align: justify;
                }}
                .print-button {{
                    background-color: #4CAF50;
                    color: white;
                    padding: 10px 20px;
                    border: none;
                    border-radius: 5px;
                    cursor: pointer;
                    font-size: 16px;
                    margin: 20px 0;
                }}
                .report-header {{ 
                    text-align: center; 
                    margin-bottom: 30px;
                }}
                .report-title {{
                    font-size: 28pt;
                    font-weight: bold;
                    margin-bottom: 20px;
                }}
                .report-content {{
                    text-align: justify;
                }}
                @page {{ 
                    size: A4; 
                    margin: 2cm;
                }}
            </style>
        </head>
        <body>
            <div class="no-print" style="text-align: center;">
                <button onclick="window.print()" class="print-button">
                    打印为PDF
                </button>
            </div>
            
            <div class="report-header">
                <div class="report-title">AI实训报告</div>
                <div>{title_safe}</div>
                <div style="margin-top:10px; color:#666;">生成时间：{datetime.now().strftime('%Y-%m-%d')}</div>
            </div>
            
            <div class="report-content">
                {content}
            </div>
            
            <div class="no-print" style="text-align: center; margin-top:30px">
                <button onclick="window.print()" class="print-button">
                    打印为PDF
                </button>
            </div>
        </body>
        </html>
        """
        
        # 返回HTML响应，直接在浏览器中显示
        return HTMLResponse(content=html_content)
        
    except Exception as e:
        logger.error(f"HTML生成错误: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"HTML导出失败: {str(e)}")

@router.get("/{report_id}/export/txt")
async def export_report_txt(
    report_id: int,
    db: Session = Depends(get_db)
):
    """导出报告为纯文本格式"""
    try:
        # 查找报告
        report = db.query(Report).filter(Report.id == report_id).first()
        if report is None:
            raise HTTPException(status_code=404, detail="报告不存在")
        
        # 准备文本内容
        text_content = f"""
{report.title}
{'=' * len(report.title)}

{report.content}

-----------------------------
生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
        """
        
        # 设置安全的文件名
        safe_filename = f"report_{report_id}.txt"
        
        # 返回文本响应
        return Response(
            content=text_content,
            media_type="text/plain",
            headers={"Content-Disposition": f"attachment; filename={safe_filename}"}
        )
        
    except Exception as e:
        logger.error(f"文本导出错误: {str(e)}")
        raise HTTPException(status_code=500, detail=f"文本导出失败: {str(e)}")

