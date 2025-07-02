from typing import List
import os
import uuid
import shutil
import logging
from typing import List, Dict, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, status
from sqlalchemy.orm import Session
from datetime import datetime
from fastapi.responses import StreamingResponse
from docx import Document
import io
import re

from ..config import settings
from ..database import get_db
from ..models.user import User
from ..models.template import Template
from ..schemas.template import TemplateCreate, TemplateResponse
from ..utils.security import get_current_user

TEMPLATE_DIR = "/home/ai_report_system/templates/docx"

router = APIRouter(prefix=f"{settings.API_V1_STR}/templates", tags=["模板"])
logger = logging.getLogger(__name__)

@router.post("/", response_model=TemplateResponse, status_code=status.HTTP_201_CREATED)
async def create_template(
    template_data: TemplateCreate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """创建新模板"""
    new_template = Template(
        name=template_data.name,
        content=template_data.content,
        user_id=current_user.id
    )
    
    db.add(new_template)
    db.commit()
    db.refresh(new_template)
    
    return new_template

@router.post("/upload-docx", status_code=status.HTTP_200_OK)
async def upload_docx_template(
    file: UploadFile = File(...),
    name: str = Form(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """上传Word格式的报告模板"""
    try:
        logger.info(f"接收到模板上传请求: 文件名={file.filename}, 模板名={name}, 用户ID={current_user.id}")
        
        # 检查文件类型
        if not file.filename.lower().endswith(".docx"):
            logger.warning(f"文件类型错误: {file.filename}")
            raise HTTPException(status_code=400, detail="只支持.docx格式文件")
        
        # 生成唯一文件名，保留原始扩展名
        original_filename = file.filename
        file_extension = original_filename.split('.')[-1]
        unique_filename = f"{uuid.uuid4().hex}_{datetime.now().strftime('%Y%m%d%H%M%S')}.{file_extension}"
        
        # 确保目录存在
        os.makedirs(TEMPLATE_DIR, exist_ok=True)
        file_path = os.path.join(TEMPLATE_DIR, unique_filename)
        
        logger.info(f"保存文件到路径: {file_path}")
        
        # 保存文件到磁盘
        try:
            # 读取上传的文件内容
            content = await file.read()
            
            # 写入到目标路径
            with open(file_path, "wb") as buffer:
                buffer.write(content)
                
            logger.info(f"文件保存成功, 大小: {len(content)} 字节")
        except Exception as save_error:
            logger.error(f"保存文件失败: {str(save_error)}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"保存文件失败: {str(save_error)}")
        
        # 创建模板记录
        try:
            new_template = Template(
                name=name,
                content=f"DOCX模板 - {original_filename}",  # 简单描述
                user_id=current_user.id,
                file_path=file_path,
                original_filename=original_filename,
                is_docx=True
            )
            
            db.add(new_template)
            db.commit()
            db.refresh(new_template)
            logger.info(f"模板记录创建成功: ID={new_template.id}")
        except Exception as db_error:
            # 如果数据库操作失败，删除已保存的文件
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"由于数据库错误，已删除文件: {file_path}")
            
            logger.error(f"数据库操作失败: {str(db_error)}", exc_info=True)
            db.rollback()
            raise HTTPException(status_code=500, detail=f"数据库操作失败: {str(db_error)}")
        
        return {
            "id": new_template.id, 
            "name": new_template.name, 
            "original_filename": new_template.original_filename,
            "message": "DOCX模板上传成功"
        }
        
    except HTTPException:
        # 重新抛出HTTP异常
        raise
    except Exception as e:
        # 记录详细错误
        logger.error(f"模板上传处理出错: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"模板上传失败: {str(e)}")


@router.get("/", response_model=List[TemplateResponse])
async def read_templates(
    skip: int = 0,
    limit: int = 100,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取当前用户的模板列表"""
    templates = db.query(Template).filter(Template.user_id == current_user.id).offset(skip).limit(limit).all()
    return templates

@router.get("/{template_id}", response_model=TemplateResponse)
async def read_template(
    template_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """获取指定模板详情"""
    template = db.query(Template).filter(Template.id == template_id, Template.user_id == current_user.id).first()
    if template is None:
        raise HTTPException(status_code=404, detail="模板不存在")
    return template

@router.post("/{template_id}/apply-template")
async def apply_template(
    template_id: int,
    report_content: Dict[str, str],
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """将报告内容套用到DOCX模板中"""
    try:
        # 获取模板
        template = db.query(Template).filter(Template.id == template_id, 
                                           Template.user_id == current_user.id).first()
        if template is None:
            raise HTTPException(status_code=404, detail="模板不存在")
        
        # 检查是否为DOCX模板
        if not template.is_docx or not template.file_path:
            raise HTTPException(status_code=400, detail="此模板不是DOCX格式或文件路径为空")
        
        # 检查文件是否存在
        file_path = template.file_path
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="模板文件不存在")
        
        logger.info(f"套用模板: ID={template_id}, 文件={file_path}")
        
        # 加载DOCX文件
        try:
            doc = Document(file_path)
            logger.info(f"成功加载DOCX模板: {file_path}")
        except Exception as e:
            logger.error(f"加载DOCX模板失败: {str(e)}")
            raise HTTPException(status_code=500, detail=f"无法加载DOCX模板: {str(e)}")
        
        # 替换所有占位符
        replacements_count = 0
        placeholders_found = set()
        
        # 1. 替换段落中的占位符
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            
            # 查找所有占位符
            matches = re.findall(r'\{([^{}]+)\}', original_text)
            
            # 记录找到的占位符
            placeholders_found.update(matches)
            
            # 替换占位符
            for placeholder in matches:
                if placeholder in report_content:
                    replacement = report_content[placeholder]
                    modified_text = modified_text.replace(f"{{{placeholder}}}", replacement)
                    replacements_count += 1
            
            # 如果有变更，更新段落文本
            if modified_text != original_text:
                para.text = modified_text
        
        # 2. 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        modified_text = original_text
                        
                        # 查找所有占位符
                        matches = re.findall(r'\{([^{}]+)\}', original_text)
                        
                        # 记录找到的占位符
                        placeholders_found.update(matches)
                        
                        # 替换占位符
                        for placeholder in matches:
                            if placeholder in report_content:
                                replacement = report_content[placeholder]
                                modified_text = modified_text.replace(f"{{{placeholder}}}", replacement)
                                replacements_count += 1
                        
                        # 如果有变更，更新段落文本
                        if modified_text != original_text:
                            para.text = modified_text
        
        logger.info(f"找到占位符: {placeholders_found}")
        logger.info(f"完成替换次数: {replacements_count}")
        
        # 保存到内存
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        file_name = f"report_{template_id}_{timestamp}.docx"
        
        # 返回生成的文件
        return StreamingResponse(
            output_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={file_name}"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"套用模板出错: {str(e)}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"套用模板失败: {str(e)}")

@router.delete("/{template_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_template(
    template_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """删除模板，如果是DOCX模板同时删除文件"""
    template = db.query(Template).filter(Template.id == template_id, 
                                        Template.user_id == current_user.id).first()
    if template is None:
        raise HTTPException(status_code=404, detail="模板不存在")
    
    # 如果是DOCX模板，删除对应的文件
    if template.is_docx and template.file_path:
        file_path = template.file_path
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"已删除模板文件: {file_path}")
            except Exception as e:
                logger.error(f"删除模板文件失败: {str(e)}")
                # 即使文件删除失败，也继续删除数据库记录
    
    # 删除数据库记录
    db.delete(template)
    db.commit()
    logger.info(f"已删除模板记录: ID={template_id}")
    
    return None